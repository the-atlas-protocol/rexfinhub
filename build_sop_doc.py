"""Concise SOP: what went wrong and why."""
from docx import Document
from docx.shared import Pt, RGBColor
import os

doc = Document()
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(10)
style.paragraph_format.space_after = Pt(4)

def heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)

def bullet(text):
    doc.add_paragraph(text, style='List Bullet')

title = doc.add_heading('REX Asia Report - What Went Wrong', level=0)
for run in title.runs:
    run.font.color.rgb = RGBColor(0, 0, 0)

heading('Root Causes')

bullet('Sean\'s database only stored global AUM for funds with Asia positions (84 of 93 REX ETPs). '
       'Total REX AUM was understated by ~$190M, inflating the "% of AUM in Asia" from 19.6% to 20.2%.')

bullet('Sean\'s global AUM values did not match Bloomberg. 259 fund-month combinations were wrong '
       '(some by 2-80x), with ETN values consistently incorrect. These fed directly into the '
       'flow/market move calculation, producing wrong splits for nearly every fund.')

bullet('Sean\'s price data had 221 mismatches vs Bloomberg. Prices drive the repricing formula '
       'for quarterly reporters (18% of total Asia AUM), meaning repriced positions were also affected.')

bullet('Seven database views computed aggregated metrics using only Asia-active funds in the denominator. '
       'Any query through these views produced deflated percentages and distorted flow estimates. '
       'These views have been quarantined.')

bullet('All derived values (flows, market move, MoM, suite KPIs, chart data) were hardcoded in the '
       'report HTML and copied between versions. When the underlying database was corrected, '
       'the report values were never recalculated.')

heading('Resolution')

bullet('Rebuilt etp_monthly_fund from Bloomberg daily file (data_aum for ETFs, microsector for ETNs, '
       'data_price for prices) for all 13 months. Zero mismatches vs Bloomberg after rebuild.')

bullet('All report values now recalculated from the corrected database. Verified via automated audit: '
       '84/84 fund AUM, flows, MoM match; 13/13 timeline months match; flow balances verified; '
       'cross-checked against Grace source data and Seamus daily email.')

out_path = os.path.join('C:/Projects/rex-asia/reports/2026-02', 'REX_Asia_SOP.docx')
try:
    os.remove(out_path)
except (FileNotFoundError, PermissionError):
    pass
doc.save(out_path)
print(f"Written: {out_path}")
