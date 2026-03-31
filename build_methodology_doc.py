"""Concise data notes doc."""
from docx import Document
from docx.shared import Pt, RGBColor
import os

doc = Document()
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(10)
style.paragraph_format.space_after = Pt(4)

def heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)

def bullet(text):
    doc.add_paragraph(text, style='List Bullet')

title = doc.add_heading('REX Asia Report - Feb 2026 Data Notes', level=0)
for run in title.runs:
    run.font.color.rgb = RGBColor(0, 0, 0)

heading('Methodology')
bullet('Asia AUM sourced from broker reports (Grace Tse) covering 16 exchanges across 7 markets')
bullet('Global AUM sourced from Bloomberg daily file (data_aum for ETFs, microsector sheet for ETNs)')
bullet('Quarterly reporters (Futu/MooMoo, Oriental Harbour) are repriced using fund-proportional scaling between reports')
bullet('Flows estimated as: Dollar Change minus Market Move, where Market Move = Prior Asia AUM x Global MoM')

heading('Key Differences vs Grace')
bullet('Report total $1,320M vs Grace $1,388M - difference is $68M from repricing quarterly carry-forwards')
bullet('All reported (non-quarterly) sources match Grace exactly')

heading('March - Quarterly Data')
bullet('Q1 2026 reports (Futu, Oriental Harbour, ViewTrade) will replace ~$240M of repriced estimates with actuals')

out_path = os.path.join('C:/Projects/rex-asia/reports/2026-02', 'REX_Asia_Data_Notes_Feb26.docx')
for f in [out_path, out_path.replace('.docx', '_v2.docx')]:
    try:
        os.remove(f)
    except (FileNotFoundError, PermissionError):
        pass

doc.save(out_path)
print(f"Written: {out_path}")
